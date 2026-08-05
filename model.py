# import pickle

# model = pickle.load(open("model.pkl", "rb"))
# vectorizer = pickle.load(open("vectorizer.pkl", "rb"))

# def predict_career(skills):

#     X = vectorizer.transform([skills])

#     prediction = model.predict(X)[0]

#     return prediction
# # import pickle

# # model = pickle.load(open("model.pkl", "rb"))
# # vectorizer = pickle.load(open("vectorizer.pkl", "rb"))

# # def predict_career(skills):
# #     skills_text = " ".join(skills.split(","))

# #     X = vectorizer.transform([skills_text])
# #     prediction = model.predict(X)[0]

# #     return prediction


# # def ats_score(resume):

# #     keywords = [
# #         "python","machine learning","data","analysis",
# #         "project","team","communication","sql"
# #     ]

# #     score = 0

# #     for word in keywords:
# #         if word in resume.lower():
# #             score += 10

# #     return min(score,100)
# # import PyPDF2

# # def ats_score(filepath):

# #     text = ""

# #     with open(filepath, "rb") as file:
# #         reader = PyPDF2.PdfReader(file)

# #         for page in reader.pages:
# #             if page.extract_text():
# #                 text += page.extract_text()

# #     text = text.lower()

# #     # simple keyword scoring
# #     keywords = ["python", "sql", "machine learning", "data","analysis","project","team","communication"]

# #     score = sum(word in text for word in keywords) * 25

# #     return min(score, 100)

# import PyPDF2

# # 🎯 JOB ROLE SKILLS DATABASE
# JOB_ROLES = {
#     "data_scientist": ["python", "machine learning", "data analysis", "pandas", "numpy", "sql"],
#     "web_developer": ["html", "css", "javascript", "react", "node", "mongodb"],
#     "software_engineer": ["java", "python", "c++", "dsa", "algorithms", "git"],
#     "data_analyst": ["excel", "sql", "power bi", "tableau", "python", "statistics"]
# }


# # 📄 Extract text from PDF
# def extract_text(filepath):

#     text = ""

#     with open(filepath, "rb") as file:
#         reader = PyPDF2.PdfReader(file)

#         for page in reader.pages:
#             if page.extract_text():
#                 text += page.extract_text()

#     return text.lower()


# # 🤖 MAIN ATS FUNCTION
# def ats_score(filepath, job_role="data_scientist"):

#     text = extract_text(filepath)

#     required_skills = JOB_ROLES.get(job_role, [])

#     matched = []
#     missing = []

#     for skill in required_skills:
#         if skill in text:
#             matched.append(skill)
#         else:
#             missing.append(skill)

#     # 🎯 SCORE CALCULATION
#     score = int((len(matched) / len(required_skills)) * 100)

#     # 💡 AI SUGGESTIONS
#     suggestions = []

#     if "projects" not in text:
#         suggestions.append("Add Projects section")

#     if "experience" not in text:
#         suggestions.append("Add Work Experience")

#     if len(missing) > 0:
#         suggestions.append("Add missing skills: " + ", ".join(missing[:3]))

#     if score < 50:
#         suggestions.append("Improve resume with more relevant keywords")

#     return {
#         "score": score,
#         "matched_skills": matched,
#         "missing_skills": missing,
#         "suggestions": suggestions
#     }

# def clean_text(text):
#     # Lowercase, remove special characters, remove extra spaces
#     text = text.lower()
#     text = re.sub(r'[^a-z0-9\s]', '', text) 
#     return text

# def predict_career(skills):
#     cleaned_skills = clean_text(skills) # Clean it first!
#     X = vectorizer.transform([cleaned_skills])
#     prediction = model.predict(X)[0]
#     return prediction





# import pickle
# import PyPDF2
# import re

# # ---------------------------------------------------------
# # 1. ML MODEL LOADING (Ensure these files are in your folder)
# # ---------------------------------------------------------
# try:
#     model = pickle.load(open("model.pkl", "rb"))
#     vectorizer = pickle.load(open("vectorizer.pkl", "rb"))
# except FileNotFoundError:
#     print("Warning: model.pkl or vectorizer.pkl not found. Career prediction will not work.")

# # ---------------------------------------------------------
# # 2. JOB ROLE SKILLS DATABASE
# # ---------------------------------------------------------
# JOB_ROLES = {
#     "data_scientist": ["python", "machine learning", "data analysis", "pandas", "numpy", "sql"],
#     "web_developer": ["html", "css", "javascript", "react", "node", "mongodb"],
#     "software_engineer": ["java", "python", "c++", "dsa", "algorithms", "git"],
#     "data_analyst": ["excel", "sql", "power bi", "tableau", "python", "statistics"]
# }

# # ---------------------------------------------------------
# # 3. TEXT PROCESSING & EXTRACTION
# # ---------------------------------------------------------
# def extract_text(filepath):
#     """Extracts and returns text from a PDF file."""
#     text = ""
#     with open(filepath, "rb") as file:
#         reader = PyPDF2.PdfReader(file)
#         for page in reader.pages:
#             extracted = page.extract_text()
#             if extracted:
#                 text += extracted
#     return text.lower()

# def clean_text(text):
#     """Cleans text by making it lowercase and removing special characters."""
#     text = text.lower()
#     text = re.sub(r'[^a-z0-9\s]', '', text) 
#     return text

# # ---------------------------------------------------------
# # 4. CAREER PREDICTION
# # ---------------------------------------------------------
# def predict_career(skills):
#     """Predicts a career based on inputted skills using the loaded ML model."""
#     cleaned_skills = clean_text(skills) 
#     X = vectorizer.transform([cleaned_skills])
#     prediction = model.predict(X)[0]
#     return prediction

# # ---------------------------------------------------------
# # 5. ATS SCORING FUNCTION
# # ---------------------------------------------------------
# def ats_score(filepath, job_role="data_scientist"):
#     """Scores a resume against specific job role requirements."""
#     text = extract_text(filepath)
#     required_skills = JOB_ROLES.get(job_role, [])
    
#     if not required_skills:
#         return {"error": "Job role not found in database."}

#     matched = []
#     missing = []

#     for skill in required_skills:
#         if skill in text:
#             matched.append(skill)
#         else:
#             missing.append(skill)

#     # 🎯 SCORE CALCULATION
#     score = int((len(matched) / len(required_skills)) * 100)

#     # 💡 AI SUGGESTIONS
#     suggestions = []
#     if "projects" not in text:
#         suggestions.append("Add a 'Projects' section.")
#     if "experience" not in text:
#         suggestions.append("Add a 'Work Experience' section.")
#     if len(missing) > 0:
#         suggestions.append("Add missing skills: " + ", ".join(missing[:3]))
#     if score < 50:
#         suggestions.append("Improve resume with more relevant keywords.")

#     return {
#         "score": score,
#         "matched_skills": matched,
#         "missing_skills": missing,
#         "suggestions": suggestions
#     }


# from extensions import db



# # Make sure you have db imported at the top of model.py
# # from extensions import db 

# class Course(db.Model):
#     __tablename__ = 'courses'

#     id = db.Column(db.Integer, primary_key=True)
#     title = db.Column(db.String(200), nullable=False)
#     slug = db.Column(db.String(200), nullable=False)
#     description = db.Column(db.Text, nullable=True)

#     # Note: I completely removed category_id, instructor, price, etc.
#     # Since they aren't in your form, they shouldn't be required in your model!




# class UserProgress(db.Model):
#     id = db.Column(db.Integer, primary_key=True)
#     user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
#     chapter_id = db.Column(db.Integer, db.ForeignKey('chapter.id'), nullable=False)
#     status = db.Column(db.String(20), default='not_started') # 'completed', 'in_progress'
    
#     # This ensures a user can only have one progress record per chapter
#     __table_args__ = (db.UniqueConstraint('user_id', 'chapter_id', name='_user_chapter_uc'),)






# import os
# import pickle
# import PyPDF2
# import docx
# import re
# from collections import Counter
# from extensions import db # Imported from your existing project structure

# # ---------------------------------------------------------
# # 1. ML MODEL LOADING (For Career Prediction)
# # ---------------------------------------------------------
# try:
#     model = pickle.load(open("model.pkl", "rb"))
#     vectorizer = pickle.load(open("vectorizer.pkl", "rb"))
# except FileNotFoundError:
#     print("Warning: model.pkl or vectorizer.pkl not found. Career prediction will not work.")

# # ---------------------------------------------------------
# # 2. JOB ROLE SKILLS DATABASE (Kept for legacy functions)
# # ---------------------------------------------------------
# # JOB_ROLES = {
# #     "data_scientist": ["python", "machine learning", "data analysis", "pandas", "numpy", "sql"],
# #     "web_developer": ["html", "css", "javascript", "react", "node", "mongodb"],
# #     "software_engineer": ["java", "python", "c++", "dsa", "algorithms", "git"],
# #     "data_analyst": ["excel", "sql", "power bi", "tableau", "python", "statistics"]
# # }

# # ---------------------------------------------------------
# # 3. TEXT PROCESSING & EXTRACTION (Updated for PDF & DOCX)
# # ---------------------------------------------------------
# def extract_text(filepath):
#     """Extracts and returns raw text from a PDF or DOCX file."""
#     text = ""
#     ext = os.path.splitext(filepath)[1].lower()
    
#     try:
#         if ext == '.pdf':
#             with open(filepath, "rb") as file:
#                 reader = PyPDF2.PdfReader(file)
#                 for page in reader.pages:
#                     extracted = page.extract_text()
#                     if extracted:
#                         text += extracted + " "
#         elif ext in ['.doc', '.docx']:
#             doc = docx.Document(filepath)
#             for para in doc.paragraphs:
#                 text += para.text + " "
#     except Exception as e:
#         print(f"Error extracting text: {e}")
        
#     return text

# def clean_text(text):
#     """Cleans text by making it lowercase and removing special characters."""
#     text = text.lower()
#     text = re.sub(r'[^a-z0-9\s]', '', text) 
#     return text

# # ---------------------------------------------------------
# # 4. CAREER PREDICTION (Untouched)
# # ---------------------------------------------------------
# def predict_career(skills):
#     """Predicts a career based on inputted skills using the loaded ML model."""
#     cleaned_skills = clean_text(skills) 
#     X = vectorizer.transform([cleaned_skills])
#     prediction = model.predict(X)[0]
#     return prediction

# # ---------------------------------------------------------
# # 5. NEW ATS SCORING ALGORITHM (4-Pillar Layout)
# # ---------------------------------------------------------
# STOP_WORDS = {'the', 'and', 'to', 'of', 'a', 'in', 'for', 'is', 'on', 'that', 'by', 'with', 'i', 'you', 'it', 'not', 'or', 'be', 'are'}

# def ats_score(filepath):
#     """Dynamically scores a resume based on structural ATS rules."""
#     raw_text = extract_text(filepath)
    
#     # If file is empty or unreadable
#     if not raw_text or len(raw_text.strip()) < 20:
#         return {"error": "Could not extract text. Ensure it is a valid text-based PDF or DOCX."}

#     text_lower = raw_text.lower()
#     words = re.findall(r'\b\w+\b', text_lower)
    
#     score = 100
#     report = {
#         "score": 0,
#         "parse_rate": [],
#         "impact": [],
#         "repetition": [],
#         "grammar": []
#     }

#     # --- Pillar 1: ATS Parse Rate ---
#     essential_headers = ['experience', 'education', 'skills']
#     found_headers = [h for h in essential_headers if h in text_lower]
#     if len(found_headers) == len(essential_headers):
#         report["parse_rate"].append("✅ Standard headers (Experience, Education, Skills) detected.")
#     else:
#         score -= 20
#         missing = set(essential_headers) - set(found_headers)
#         report["parse_rate"].append(f"⚠️ Missing critical headers: {', '.join(missing).title()}")

#     # --- Pillar 2: Quantifying Impact ---
#     metrics = re.findall(r'\d+%|\$\d+|[\d,]+', raw_text)
#     if len(metrics) >= 3:
#         report["impact"].append(f"✅ Found {len(metrics)} quantifiable metrics in your experience.")
#     else:
#         score -= 20
#         report["impact"].append("❌ Low use of metrics. Add numbers to show results (e.g., 'Increased sales by 20%').")

#     # --- Pillar 3: Repetition ---
#     meaningful_words = [w for w in words if w not in STOP_WORDS and len(w) > 2]
#     counts = Counter(meaningful_words)
#     overused = [word for word, count in counts.items() if count > 4]
#     if overused:
#         score -= 10
#         report["repetition"].append(f"⚠️ Overused words detected: {', '.join(overused[:3])}. Use synonyms.")
#     else:
#         report["repetition"].append("✅ Good vocabulary variety detected.")

#     # --- Pillar 4: Grammar & Formatting ---
#     if "  " in raw_text:
#         score -= 10
#         report["grammar"].append("⚠️ Multiple spacing errors found. This can confuse ATS parsers.")
#     else:
#         report["grammar"].append("✅ Clean formatting and spacing detected.")

#     report["score"] = max(0, score)
#     return report

# # ---------------------------------------------------------
# # 6. DATABASE MODELS (Untouched)
# # ---------------------------------------------------------
# class Course(db.Model):
#     __tablename__ = 'courses'
#     id = db.Column(db.Integer, primary_key=True)
#     title = db.Column(db.String(200), nullable=False)
#     slug = db.Column(db.String(200), nullable=False)
#     description = db.Column(db.Text, nullable=True)

# class UserProgress(db.Model):
#     id = db.Column(db.Integer, primary_key=True)
#     user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
#     chapter_id = db.Column(db.Integer, db.ForeignKey('chapter.id'), nullable=False)
#     status = db.Column(db.String(20), default='not_started') # 'completed', 'in_progress'
    
#     __table_args__ = (db.UniqueConstraint('user_id', 'chapter_id', name='_user_chapter_uc'),)


import os
import pickle
import PyPDF2
import docx
import re
from collections import Counter
from extensions import db  # Ensure this matches your project structure

# ---------------------------------------------------------
# 1. ML MODEL LOADING (For Career Prediction)
# ---------------------------------------------------------
try:
    model = pickle.load(open("model.pkl", "rb"))
    vectorizer = pickle.load(open("vectorizer.pkl", "rb"))
except FileNotFoundError:
    print("Warning: model.pkl or vectorizer.pkl not found. Career prediction will not work.")

# ---------------------------------------------------------
# 2. ENHANCV-TIER DICTIONARIES & CONSTANTS
# ---------------------------------------------------------
STOP_WORDS = {'the', 'and', 'to', 'of', 'a', 'in', 'for', 'is', 'on', 'that', 'by', 'with', 'it', 'or', 'be', 'are', 'at', 'as', 'this'}
WEAK_PHRASES = ['responsible for', 'duties included', 'helped with', 'worked on', 'assisted in', 'tasked with', 'handled']
CLICHES = ['team player', 'hard worker', 'think outside the box', 'synergy', 'go-getter', 'detail-oriented', 'dynamic', 'results-driven']
PRONOUNS = [r'\bi\b', r'\bme\b', r'\bmy\b', r'\bwe\b', r'\bour\b'] # Resumes should never use first-person pronouns

# ---------------------------------------------------------
# 3. TEXT PROCESSING & EXTRACTION
# ---------------------------------------------------------
def extract_text(filepath):
    """Extracts raw text from a PDF or DOCX file."""
    text = ""
    ext = os.path.splitext(filepath)[1].lower()
    try:
        if ext == '.pdf':
            with open(filepath, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
        elif ext in ['.doc', '.docx']:
            doc = docx.Document(filepath)
            for para in doc.paragraphs:
                text += para.text + "\n"
    except Exception as e:
        print(f"Error extracting text: {e}")
    return text

def clean_text(text):
    text = text.lower()
    text = re.sub(r'[^a-z0-9\s]', '', text) 
    return text

# ---------------------------------------------------------
# 4. CAREER PREDICTION (Untouched)
# ---------------------------------------------------------
def predict_career(skills):
    """Predicts a career based on inputted skills using the loaded ML model."""
    cleaned_skills = clean_text(skills) 
    X = vectorizer.transform([cleaned_skills])
    prediction = model.predict(X)[0]
    return prediction

# ---------------------------------------------------------
# 5. THE ENHANCV-CLONE ATS ENGINE
import re
from collections import Counter

# Expanded dictionaries for advanced heuristics
STRONG_ACTION_VERBS = {'developed', 'managed', 'led', 'spearheaded', 'architected', 'optimized', 'designed', 'implemented', 'orchestrated', 'streamlined', 'resolved', 'analyzed', 'facilitated'}
WEAK_PHRASES = {'responsible for', 'helped with', 'worked on', 'duties included', 'assisted in'}
PRONOUNS = {r'\bi\b', r'\bme\b', r'\bmy\b', r'\bmine\b', r'\bwe\b', r'\bous\b'}
STOP_WORDS = {'the', 'and', 'a', 'to', 'of', 'in', 'for', 'is', 'on', 'that', 'by', 'this', 'with', 'i', 'you', 'it'}

def ats_score(filepath):
    """
    Enterprise-grade ATS parsing logic.
    Analyzes structural consistency, granular impact, and semantic tone.
    """
    # 1. FIX: Actually extract the text from the file path
    raw_text = extract_text(filepath)
    
    # 2. FIX: Handle Image-based PDFs
    if not raw_text or len(raw_text.strip()) < 20:
        return {"error": "No text detected. If this is a scanned image or picture, ATS systems cannot read it. Please upload a standard text-based PDF."}

    text_lower = raw_text.lower()
    words = re.findall(r'\b[a-z]+\b', text_lower)
    word_count = len(words)
    
    # Smarter sentence/bullet extraction
    cleaned_text = re.sub(r'[•·\-\*]', '.', raw_text)
    sentences = [s.strip() for s in re.split(r'(?<=[.!?]) +|\n+', cleaned_text) if len(s.strip()) > 15]
    
    score = 100
    # FIX: Updated keys to match your frontend JavaScript (parse_rate, grammar)
    report = {"score": 0, "parse_rate": [], "impact": [], "grammar": [], "repetition": []}

    # ==========================================
    # 1. PARSE RATE & STRUCTURE 
    # ==========================================
    # Calculate a numeric Parse Rate based on standard headers found
    expected_sections = ['experience', 'education', 'skills']
    sections_found = sum(1 for sec in expected_sections if sec in text_lower)
    numeric_parse_rate = int((sections_found / len(expected_sections)) * 100)
    report["parse_rate"] = numeric_parse_rate # Sends the number to your JS circle animation!

    # Date Consistency Check
    numeric_dates = len(re.findall(r'\b\d{1,2}/\d{2,4}\b', raw_text))
    text_dates = len(re.findall(r'\b(?:january|february|march|april|may|june|july|august|september|october|november|december|jan|feb|mar|apr|jun|jul|aug|sep|oct|nov|dec)\s+\d{4}\b', text_lower))
    
    if numeric_dates > 0 and text_dates > 0:
        score -= 5
        report["grammar"].append("⚠️ <b>Date Formatting:</b> Inconsistent date formats detected. Pick one style (e.g., 'MM/YYYY') for reliable parsing.")
    elif numeric_dates == 0 and text_dates == 0:
        score -= 5
        report["grammar"].append("⚠️ <b>Timelines:</b> Few standard date formats detected. Ensure experience includes clear dates.")

    # ==========================================
    # 2. GRANULAR IMPACT ANALYSIS
    # ==========================================
    financial_metrics = re.findall(r'\$\d+(?:,\d{3})*(?:\.\d{2})?|\b\d+\s*(?:k|m|b|billion|million)\b', text_lower)
    scale_metrics = re.findall(r'\b\d+(?:\.\d+)?%', raw_text)
    time_metrics = re.findall(r'\b\d+\s+(?:hours|days|weeks|months|years)\b', text_lower)
    
    total_metrics = len(financial_metrics) + len(scale_metrics) + len(time_metrics)
    
    if total_metrics >= 5:
        report["impact"].append(f"✅ <b>High Impact:</b> Excellent quantification. Found {total_metrics} data metrics.")
    elif total_metrics >= 2:
        score -= 5
        report["impact"].append(f"⚠️ <b>Moderate Impact:</b> Found {total_metrics} data points. Try adding more percentages or financial values.")
    else:
        score -= 15
        report["impact"].append(" <b>Low Impact:</b> Lacking measurable results. Add numbers to prove your scale and efficiency.")

    # ==========================================
    # 3. SEMANTIC STYLE & ACTION VERBS (GRAMMAR)
    # ==========================================
    body_sentences = sentences[3:] if len(sentences) > 3 else sentences
    weak_start_count = 0
    
    for sentence in body_sentences:
        first_word = sentence.split()[0].lower()
        first_word = re.sub(r'[^a-z]', '', first_word)
        if first_word in ['i', 'we', 'the', 'a', 'this', 'responsible', 'helped']:
            weak_start_count += 1

    if weak_start_count > 2:
        score -= 10
        report["grammar"].append(f" <b>Weak Bullet Starts:</b> {weak_start_count} sentences start with weak words. Use action verbs (e.g., 'Developed').")

    found_weak = False
    for sentence in body_sentences:
        for phrase in WEAK_PHRASES:
            if phrase in sentence.lower():
                score -= 5
                snippet = sentence[:40] + "..."
                report["grammar"].append(f" <b>Weak Phrasing:</b> Replace '<i>{phrase}</i>' with definitive verbs. (Ex: <i>\"{snippet}\"</i>)")
                found_weak = True
                break
        if found_weak: break

    # ==========================================
    # 4. REPETITION & KEYWORD DENSITY
    # ==========================================
    seen_sentences = set()
    duplicate_issues = []
    
    for sentence in body_sentences:
        normalized_sentence = sentence.strip().lower()
        words_in_sentence = normalized_sentence.split()
        
        if len(words_in_sentence) < 8 or len(normalized_sentence) < 45: 
            continue
            
        if normalized_sentence in seen_sentences:
            duplicate_issues.append(sentence.strip().rstrip('.!?'))
        else:
            seen_sentences.add(normalized_sentence)

    if duplicate_issues:
        score -= 15 
        snippet = duplicate_issues[0][:50] + "..."
        report["repetition"].append(f" <b>Duplicate Bullets:</b> You have copy-pasted exact descriptions. <i>\"{snippet}\"</i>")

    # Keyword Density
    if word_count > 0:
        meaningful_words = [w for w in words if w not in STOP_WORDS and len(w) > 3]
        counts = Counter(meaningful_words)
        keyword_stuffing_threshold = word_count * 0.04 
        stuffed_words = [(word, count) for word, count in counts.items() if count > keyword_stuffing_threshold and count > 5]
        
        if stuffed_words:
            score -= 10
            word = stuffed_words[0][0]
            report["repetition"].append(f"⚠️ <b>Keyword Density:</b> \"<b>{word.title()}</b>\" is used {stuffed_words[0][1]} times. Use synonyms.")

    report["score"] = max(0, min(100, score))
    return report
# ---------------------------------------------------------
# 6. DATABASE MODELS
# ---------------------------------------------------------
class Course(db.Model):
    __tablename__ = 'courses'
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(200), nullable=False)
    slug = db.Column(db.String(200), nullable=False)
    description = db.Column(db.Text, nullable=True)

class UserProgress(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    chapter_id = db.Column(db.Integer, db.ForeignKey('chapter.id'), nullable=False)
    status = db.Column(db.String(20), default='not_started')
    
    __table_args__ = (db.UniqueConstraint('user_id', 'chapter_id', name='_user_chapter_uc'),)

from extensions import db
from datetime import datetime

class AtsLog(db.Model):
    __tablename__ = 'ats_logs'  # This links it to the table in your DB Browser

    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(255), nullable=True)
    score = db.Column(db.Integer, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    user_id = db.Column(db.Integer, nullable=True)



class User(db.Model):
    __tablename__ = 'users'
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(150), unique=True, nullable=False)
    email = db.Column(db.String(150), unique=True, nullable=True)
    password = db.Column(db.String(200), nullable=False)
    mobile = db.Column(db.String(20), nullable=True)  # Changed from role to mobile
    resume = db.Column(db.String(500), nullable=True)
    profile_pic = db.Column(db.String(500), nullable=True)  # Already exists
    profile = db.relationship('Profile', backref='user', uselist=False)

class Profile(db.Model):
    __tablename__ = 'profiles'
    
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('users.id', ondelete='CASCADE'), nullable=False, unique=True)
    
    # Personal Info
    about = db.Column(db.Text, nullable=True)
    
    # Professional Info
    work_experience = db.Column(db.Text, nullable=True)
    education = db.Column(db.Text, nullable=True)
    certifications = db.Column(db.Text, nullable=True)
    
    # Skills & Languages (separated)
    skills = db.Column(db.Text, nullable=True)      # Technical skills only
    languages = db.Column(db.Text, nullable=True)   # Languages known (English, Hindi, etc.)
    interests = db.Column(db.Text, nullable=True)
    
    # Projects & Achievements
    projects = db.Column(db.Text, nullable=True)
    achievements = db.Column(db.Text, nullable=True)
    
    # Social Links
    linkedin_url = db.Column(db.String(255), nullable=True)
    github_url = db.Column(db.String(255), nullable=True)
   
    